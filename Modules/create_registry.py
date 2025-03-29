import os
import re
import openpyxl
import tkinter as tk
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tkinter import messagebox, ttk
from PyPDF2 import PdfReader, PdfWriter
from docx2pdf import convert
from Modules.utils import get_cell_value, clear_ui, select_excel_file, format_date

REGISTER_TITLE_TEMPLATE = "Templates/Word_templates/Register_title_template.docx"
REGISTER_TABLE_TEMPLATE = "Templates/Word_templates/Register_table_template.docx"
CERTIFICATES = "Templates/Certificates/"

MAIN_DATA = [
    {'name': 'OBJECTNAME', 'cell_row': '1', 'cell_column': 'B', 'description_cell_column': 'A'},
    {'name': 'CUSTOMER', 'cell_row': '14', 'cell_column': 'B', 'description_cell_column': 'A'},
    {'name': 'CONTRACTOR', 'cell_row': '15', 'cell_column': 'B', 'description_cell_column': 'A'},
    {'name': 'DESIGNORGANISATION', 'cell_row': '16', 'cell_column': 'B', 'description_cell_column': 'A'}
]

SUBOBJECT_DATA = [
    {'name': 'SUBOBJECTNAME', 'cell_column': 'B', 'cell_row': '1'},
    {'name': 'ACTNUMBER', 'cell_column': 'A'},
    {'name': 'EXECUTIONDATE', 'cell_column': 'B'},
    {'name': 'WORKNAMING', 'cell_column': 'C'},
    {'name': 'ALBUMNAME', 'cell_column': 'D'},
    {'name': 'PAGE', 'cell_column': 'E'},
    {'name': 'MATERIALS', 'cell_column': 'F'},
    {'name': 'EXECUTIVEDIAGRAM', 'cell_column': 'G'},
    {'name': 'LABORATORY', 'cell_column': 'H'},
    {'name': 'ENDDATE', 'cell_column': 'I'},
    {'name': 'NEXTWORKS', 'cell_column': 'J'}
]

class CreateRegistry:
    def __init__(self, root, menu_button, menu):
        self.root = root
        self.menu_button = menu_button
        self.menu = menu
        self.wb = None
        self.excel_file_name = None
        self.excel_file_path = None
        self.entries = {}
        self.sheet_vars = {}
        self.all_var = None

    def load_ui(self, project_name=None):
        wb, excel_file_name, excel_file_path = select_excel_file(project_name)
        if wb:
            self.wb = wb
            self.excel_file_name = excel_file_name
            self.excel_file_path = excel_file_path
            self.show_acts_ui()

    def show_acts_ui(self):
        clear_ui(self.root, self.menu_button, self.menu)
        main_sheet = self.wb['Main data']
        
        self.entries.clear()
        for i, value_data in enumerate(MAIN_DATA):
            label_text = get_cell_value(main_sheet, value_data['cell_row'], value_data['description_cell_column'])
            value = get_cell_value(main_sheet, value_data['cell_row'], value_data['cell_column'])
            tk.Label(self.root, text=label_text).grid(row=i, column=0, padx=5, pady=5, sticky="e")
            self.entries[value_data['name']] = tk.Entry(self.root, width=100)
            self.entries[value_data['name']].grid(row=i, column=1, padx=5, pady=5)
            self.entries[value_data['name']].insert(0, value)

        frame = tk.Frame(self.root)
        frame.grid(row=11, column=0, columnspan=2, padx=5, pady=5, sticky="nsew")
        canvas = tk.Canvas(frame)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        tk.Label(scrollable_frame, text="Выберите листы:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.sheet_vars.clear()
        for i, sheet_name in enumerate(self.wb.sheetnames, start=1):
            if sheet_name not in {'Main data', 'Contents'} and not re.match(r'^[!_]', sheet_name):
                var = tk.BooleanVar()
                tk.Checkbutton(scrollable_frame, text=sheet_name, variable=var).grid(row=i, column=0, padx=5, pady=2, sticky="w")
                self.sheet_vars[sheet_name] = var
        canvas.grid(row=1, column=0, sticky="nsew")
        scrollbar.grid(row=1, column=1, sticky="ns")
        frame.grid_rowconfigure(1, weight=1)
        frame.grid_columnconfigure(0, weight=1)
        self.all_var = tk.BooleanVar()
        tk.Checkbutton(frame, text="Выбрать все", variable=self.all_var, command=self.toggle_all_sheets).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        generate_button = tk.Button(self.root, text="Создать реестр", command=self.generate_document)
        generate_button.grid(row=12, column=0, columnspan=2, pady=10)

    def save_table_to_excel(self, table_data, sheet_dir, subobject_name):
        try:
            # Создаем новый Excel файл
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Реестр"
            
            # Заголовки таблицы
            headers = ["№ п/п", "Наименование", "Страницы"]
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            # Заполняем данными
            for row_idx, data in table_data.items():
                ws.cell(row=row_idx + 1, column=1, value=data['idx'])
                ws.cell(row=row_idx + 1, column=2, value=data['content'])
                ws.cell(row=row_idx + 1, column=3, value=data['page'])
            
            # Форматирование
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Путь для сохранения - тот же что у PDF
            excel_path = os.path.join(sheet_dir, f"Реестр {subobject_name}.xlsx")
            wb.save(excel_path)
            return excel_path
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить Excel файл: {str(e)}")
            return None

    def save_to_excel(self):
        if self.wb is None:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл Excel!")
            return
        try:
            main_sheet = self.wb['Main data']
            for value_data in MAIN_DATA:
                cell = f"{value_data['cell_column']}{value_data['cell_row']}"
                main_sheet[cell].value = self.entries[value_data['name']].get()
            try:
                self.wb.save(self.excel_file_path)
                messagebox.showinfo("Успех", f"Данные сохранены в {self.excel_file_path}")
            except PermissionError:
                messagebox.showerror("Ошибка", "Нет прав доступа для записи в файл Excel.")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке данных: {str(e)}")

    def generate_document(self):
        if self.wb is None:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл Excel!")
            return
        selected_sheets = [sheet_name for sheet_name, var in self.sheet_vars.items() if var.get()]
        if not selected_sheets:
            messagebox.showwarning("Предупреждение", "Выберите хотя бы один лист!")
            return
        self.save_to_excel()
        created_files = []
        errors = []
        base_dir = os.path.join(os.getcwd(), "Documents")
        os.makedirs(base_dir, exist_ok=True)
        excel_dir = os.path.join(base_dir, self.excel_file_name)
        os.makedirs(excel_dir, exist_ok=True)
        for sheet_name in selected_sheets:
            ws = self.wb[sheet_name]
            last_row = ws.max_row
            while last_row > 0 and all(cell.value is None for cell in ws[last_row]):
                last_row -= 1
            sheet_dir = os.path.join(excel_dir, sheet_name)
            os.makedirs(sheet_dir, exist_ok=True)
            output_path = None
            try:
                output_path, doc = self.generate_document_title(ws, last_row, sheet_dir)
                self.doc = doc
                self.output_path = output_path
                doc, document_entries = self.generate_document_table(doc, ws, last_row, sheet_dir)
                
                # Получаем данные таблицы и SUBOBJECTNAME для имени файла
                table_data, _, _ = fill_table_data(ws, last_row, sheet_dir)
                subobject_name = get_cell_value(ws, SUBOBJECT_DATA[0]['cell_row'], SUBOBJECT_DATA[0]['cell_column'])
                
                # Сохраняем таблицу в Excel
                excel_path = self.save_table_to_excel(table_data, sheet_dir, subobject_name)
                if excel_path:
                    created_files.append(excel_path)
                
                final_pdf_path = self.generate_document_content(document_entries, sheet_dir)
                created_files.append(final_pdf_path)
            except Exception as e:
                error_path = output_path if output_path else f"для листа {sheet_name}"
                errors.append(f"Ошибка при создании {error_path}: {str(e)}")
        if created_files:
            messagebox.showinfo("Успех", f"Созданы файлы:\n" + "\n".join(created_files))
        if errors:
            messagebox.showerror("Ошибки", "\n".join(errors))

    def generate_document_title(self, ws, last_row, sheet_dir):
        replacements = {}
        for value_data in MAIN_DATA:
            replacements[value_data['name']] = self.entries[value_data['name']].get()
        for obj in SUBOBJECT_DATA:
            if obj['name'] == 'SUBOBJECTNAME':
                replacements[obj['name']] = get_cell_value(ws, obj['cell_row'], obj['cell_column'])
            elif obj['name'] == 'EXECUTIONDATE':
                replacements[obj['name']] = get_cell_value(ws, 3, obj['cell_column'])
            elif obj['name'] == 'ENDDATE':
                replacements[obj['name']] = get_cell_value(ws, last_row, obj['cell_column'])
        output_path = os.path.join(sheet_dir, f"Реестр {replacements['SUBOBJECTNAME']}.docx")
        doc = self.create_word_doc(REGISTER_TITLE_TEMPLATE, replacements)
        return output_path, doc        

    def generate_document_table(self, doc, ws, last_row, sheet_dir):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        try:
            table_doc = Document(REGISTER_TABLE_TEMPLATE)
            for table in table_doc.tables:
                table_data, data_rows_needed, document_entries = fill_table_data(ws, last_row, sheet_dir)
                header_rows = 1
                new_table = doc.add_table(rows=header_rows + data_rows_needed, cols=len(table.columns))
                copy_header_values(new_table, table)
                fill_table_with_data(new_table, table_data, header_rows)
                table_paragraph = new_table._element.getparent()
                if table_paragraph.tag.endswith('p'):
                    table_paragraph.paragraph_format.left_indent = 250000
                    table_paragraph.paragraph_format.right_indent = 250000
                new_table.autofit = True
            return doc, document_entries
        except Exception as e:
            raise Exception(f"Ошибка в generate_document_table: {str(e)}")

    def generate_document_content(self, document_entries, sheet_dir):
        try:
            # Сохраняем DOCX с титульным листом и таблицей
            temp_docx_path = self.output_path
            self.doc.save(temp_docx_path)
            
            # Конвертируем DOCX в PDF
            temp_pdf_path = os.path.splitext(temp_docx_path)[0] + "_temp.pdf"
            convert(temp_docx_path, temp_pdf_path)
            
            # Создаем итоговый PDF с нормализацией размера страниц (A4: 595 x 842 pt)
            output_pdf_path = os.path.splitext(temp_docx_path)[0] + ".pdf"
            pdf_writer = PdfWriter()
            standard_width, standard_height = 595, 842  # A4 в пунктах

            # Добавляем страницы из временного PDF (титульный лист и таблица)
            with open(temp_pdf_path, 'rb') as temp_pdf_file:
                temp_pdf_reader = PdfReader(temp_pdf_file)
                for page in temp_pdf_reader.pages:
                    page.scale_to(standard_width, standard_height)
                    pdf_writer.add_page(page)

            # Добавляем документы по порядку из document_entries
            acts_dir = os.path.join(sheet_dir, "Акты")
            os.makedirs(acts_dir, exist_ok=True)
            if not document_entries:
                messagebox.showwarning("Предупреждение", "Список документов пуст. Дополнительное содержимое не будет добавлено.")
            else:
                for entry in document_entries:
                    # Добавляем акт
                    act_path = entry.get('act_path')
                    if act_path:
                        docx_path = os.path.join(acts_dir, f"{act_path}.docx")
                        if os.path.isfile(docx_path):
                            try:
                                act_pdf_path = os.path.join(sheet_dir, f"{act_path}_temp.pdf")
                                convert(docx_path, act_pdf_path)
                                with open(act_pdf_path, 'rb') as act_pdf_file:
                                    act_pdf_reader = PdfReader(act_pdf_file)
                                    for page in act_pdf_reader.pages:
                                        page.scale_to(standard_width, standard_height)
                                        pdf_writer.add_page(page)
                                os.remove(act_pdf_path)
                            except Exception as e:
                                messagebox.showerror("Ошибка", f"Ошибка при добавлении акта {act_path}: {str(e)}")
                        else:
                            messagebox.showwarning("Предупреждение", f"Акт {act_path}.docx не найден в папке 'Акты'.")

                    # Добавляем остальные документы (материалы, схемы, протоколы)
                    for pdf_path in entry.get('pdf_paths', []):
                        if pdf_path and os.path.isfile(pdf_path):
                            try:
                                with open(pdf_path, 'rb') as pdf_file:
                                    pdf_reader = PdfReader(pdf_file)
                                    for page in pdf_reader.pages:
                                        page.scale_to(standard_width, standard_height)
                                        pdf_writer.add_page(page)
                            except Exception as e:
                                messagebox.showerror("Ошибка", f"Ошибка при добавлении {pdf_path}: {str(e)}")

            # Сохраняем итоговый PDF
            with open(output_pdf_path, 'wb') as output_file:
                pdf_writer.write(output_file)

            # Удаляем временные файлы
            os.remove(temp_docx_path)
            os.remove(temp_pdf_path)

            return output_pdf_path
        
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при создании итогового PDF: {str(e)}")
            return None

    def create_word_doc(self, template_path, replacements):
        doc = Document(template_path)
        for para in doc.paragraphs:
            for run in para.runs:
                for key, value in replacements.items():
                    if key == run.text:
                        run.text = run.text.replace(key, value)
        return doc

    def toggle_all_sheets(self):
        select_all = self.all_var.get()
        for var in self.sheet_vars.values():
            var.set(select_all)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tcPr.append(borders)

def add_table_row_data(table_data, row_idx, content, page):
    table_data[row_idx] = {
        'idx': str(row_idx),
        'content': content,
        'page': page
    }
    return row_idx + 1

def process_split_text(table_data, row_idx, text, page, file_path):
    pdf_paths = []
    if ';' in text:
        parts = [part.strip() for part in text.split(';') if part.strip()]
        for part in parts:
            pdf_path, plus_pages = find_pdf_and_count_pages(file_path, part)
            if pdf_path:
                pdf_paths.append(pdf_path)
            if plus_pages is not None:
                page += plus_pages
            pages_add = add_pages(page, plus_pages if plus_pages is not None else 0)
            row_idx = add_table_row_data(table_data, row_idx, part, pages_add)
    elif text != '':
        pdf_path, plus_pages = find_pdf_and_count_pages(file_path, text)
        if pdf_path:
            pdf_paths.append(pdf_path)
        if plus_pages is not None:
            page += plus_pages
        pages_add = add_pages(page, plus_pages if plus_pages is not None else 0)
        row_idx = add_table_row_data(table_data, row_idx, text, pages_add)
    return row_idx, page, pdf_paths

def fill_table_data(ws, last_row, sheet_dir):
    row_idx = 1
    page = 0
    table_data = {}
    document_entries = []  # Список записей для каждой строки таблицы
    
    for i in range(3, last_row + 1):
        entry = {'act_path': None, 'pdf_paths': []}
        
        # Акт
        act_number = get_cell_value(ws, i, SUBOBJECT_DATA[1]['cell_column']) or ""
        if act_number:
            act_number_cleaned = act_number.replace('/', '_')
            entry['act_path'] = act_number_cleaned
        
        # Формируем строку таблицы
        work_naming = get_cell_value(ws, i, SUBOBJECT_DATA[3]['cell_column']) or ""
        exec_date = get_cell_value(ws, i, SUBOBJECT_DATA[2]['cell_column']) or ""
        content = f"Акт скрытых работ № {act_number} {work_naming} {format_date(exec_date) if exec_date else ''}"
        page += 2
        pages_add = add_pages(page, 2)
        row_idx = add_table_row_data(table_data, row_idx, content, pages_add)

        # Материалы
        material = get_cell_value(ws, i, SUBOBJECT_DATA[6]['cell_column']) or ""
        row_idx, page, material_paths = process_split_text(table_data, row_idx, material, page, CERTIFICATES)
        entry['pdf_paths'].extend(material_paths)

        # Исполнительная схема
        schema = get_cell_value(ws, i, SUBOBJECT_DATA[7]['cell_column']) or ""
        schema_path = os.path.join(sheet_dir, "Исполнительная схема")
        os.makedirs(schema_path, exist_ok=True)
        row_idx, page, schema_paths = process_split_text(table_data, row_idx, schema, page, schema_path)
        entry['pdf_paths'].extend(schema_paths)

        # Протокол лаборатории
        laboratory = get_cell_value(ws, i, SUBOBJECT_DATA[8]['cell_column']) or ""
        laboratory_path = os.path.join(sheet_dir, "протокол")
        os.makedirs(laboratory_path, exist_ok=True)
        row_idx, page, lab_paths = process_split_text(table_data, row_idx, laboratory, page, laboratory_path)
        entry['pdf_paths'].extend(lab_paths)

        # Добавляем запись в список
        document_entries.append(entry)

    return table_data, row_idx - 1, document_entries

def fill_table_with_data(new_table, table_data, header_rows):
    for row_idx in range(1, len(table_data) + 1):
        table_row_idx = header_rows + row_idx - 1
        cells = new_table.rows[table_row_idx].cells
        cells[0].text = table_data[row_idx]['idx']
        cells[1].text = table_data[row_idx]['content']
        cells[2].text = str(table_data[row_idx]['page'])
        for cell in cells:
            set_cell_borders(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = 0
                para.paragraph_format.space_after = 0
                para.paragraph_format.line_spacing = 1.0

def copy_header_values(new_table, template_table):
    header_row = template_table.rows[0]
    new_header_row = new_table.rows[0]
    for j, cell in enumerate(header_row.cells):
        new_cell = new_header_row.cells[j]
        new_cell.text = cell.text
        set_cell_borders(new_cell)
        for para in new_cell.paragraphs:
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
            para.paragraph_format.line_spacing = 1.0

def find_pdf_and_count_pages(folder_path, pdf_name):
    try:
        if not os.path.isdir(folder_path):
            messagebox.showerror("Ошибка", f"Папка '{folder_path}' не существует.")
            return None, None
        if '/' in pdf_name:
            pdf_name = pdf_name.replace('/', '_')
        if '\r' in pdf_name or '\n' in pdf_name:
            pdf_name = pdf_name.replace('\r\n', ' ').replace('\n', ' ').replace('\r', ' ')
        if '№' in pdf_name and 'Протокол лабораторных испытаний' not in pdf_name and 'Исполнительная схема' not in pdf_name:
            start_idx = pdf_name.find('№')
            if start_idx != -1:
                text_after_num = pdf_name[start_idx:].strip()
                pdf_name = text_after_num.split(' ', 1)[0]
        if not pdf_name.lower().endswith('.pdf'):
            pdf_name_with_ext = pdf_name + '.pdf'
        else:
            pdf_name_with_ext = pdf_name
        pdf_path = os.path.join(folder_path, pdf_name_with_ext)
        if not os.path.isfile(pdf_path):
            messagebox.showerror("Ошибка", f"Файл '{pdf_name_with_ext}' не найден.")
            return None, None
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
        return pdf_path, num_pages if num_pages else 0
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при обработке PDF: {str(e)}")
        return None, None
    
def add_pages(page, quantity):
    if quantity < 2:
        pages_add = f"{page}"
    elif quantity == 2:
        pages_add = f"{page - 1}, {page}"
    elif quantity > 2:
        pages_add = f"{page - (quantity - 1)} - {page}"
    return pages_add