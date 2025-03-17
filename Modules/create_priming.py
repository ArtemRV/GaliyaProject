from docx import Document
import tkinter as tk
from tkinter import messagebox
import random
import os
from Modules.utils import format_date, get_cell_value, clear_ui, select_excel_file

word_file = "Templates/Word_templates/Priming_template.docx"

main_data = [
    {'name': 'NUMBER', 'cell_column': 'A'},
    {'name': 'EXECUTIONISSUEDATEMONTH', 'cell_column': 'B'},
    {'name': 'EXECUTIONISSUEDATE', 'cell_column': 'B'},
    {'name': 'EXECUTIONDATEWORK', 'cell_column': 'C'},
    {'name': 'OBJECT', 'cell_column': 'D'},
    {'name': 'SAMPLE', 'cell_column': 'E'}
]

class CreatePriming:
    def __init__(self, root, menu_button, menu):
        self.root = root
        self.menu_button = menu_button
        self.menu = menu
        self.wb = None
        self.excel_file_name = None
        self.excel_file_path = None
        self.entries = {}
        self.sheet_vars = {}
        self.all_var = None

    def load_ui(self, project_name=None):
        wb, excel_file_name, excel_file_path = select_excel_file(project_name)

        if wb:
            self.wb = wb
            self.excel_file_name = excel_file_name
            self.excel_file_path = excel_file_path
            self.show_passport_ui()

    def show_passport_ui(self):
        clear_ui(self.root, self.menu_button, self.menu)

        self.entries.clear()

        generate_button = tk.Button(self.root, text="Создать лабораторию на грунт", command=self.generate_document)
        generate_button.grid(row=12, column=0, columnspan=2, pady=10)

    def generate_document(self):
        if self.wb is None:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл Excel!")
            return

        created_files = []
        errors = []

        base_dir = os.path.join(os.getcwd(), "Documents", "sppd", "priming")
        os.makedirs(base_dir, exist_ok=True)

        ws = self.wb['priming']
        last_row = ws.max_row
        while last_row > 0 and all(cell.value is None for cell in ws[last_row]):
            last_row -= 1

        for row_number in range(2, last_row + 1):
            replacements = {}

            for i in range(1, 6):
                changeres = round(random.uniform(2.048, 2.112), 3)
                changek = round((changeres / 2.167), 2)
                replacements[f"CHANGERES{i}"] = str(changeres).replace('.', ',')
                replacements[f"CHANGEK{i}"] = str(changek).replace('.', ',')

            for obj in main_data:
                if obj['name'] == 'EXECUTIONISSUEDATEMONTH':
                    date_value = get_cell_value(ws, row_number, obj['cell_column'])
                    new_value = format_date(date_value) if date_value else ""
                    replacements[obj['name']] = new_value
                else:
                    replacements[obj['name']] = get_cell_value(ws, row_number, obj['cell_column'])

                if obj['name'] == 'NUMBER':
                    priming_number = get_cell_value(ws, row_number, obj['cell_column'])
                    output_path = os.path.join(base_dir, f"Грунт_{priming_number.replace('/', '_')}.docx")

            try:
                self.create_word_doc(word_file, output_path, replacements)
                created_files.append(output_path)
            except Exception as e:
                errors.append(f"Ошибка при создании {output_path}: {str(e)}")

        if created_files:
            success_message = "Успешно созданы документы:\n" + "\n".join(created_files)
            messagebox.showinfo("Успех", success_message)
        if errors:
            error_message = "Ошибки при создании документов:\n" + "\n".join(errors)
            messagebox.showerror("Ошибка", error_message)
        if not created_files and not errors:
            messagebox.showinfo("Информация", "Документы не были созданы.")

    def create_word_doc(self, template_path, output_path, replacements):
        doc = Document(template_path)

        # Обработка параграфов в основном содержимом
        for para in doc.paragraphs:
            for run in para.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))

        # Сохранение изменённого документа
        doc.save(output_path)

    def toggle_all_sheets(self):
        select_all = self.all_var.get()
        for var in self.sheet_vars.values():
            var.set(select_all)